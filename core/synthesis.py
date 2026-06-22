from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from config import GROQ_API_KEY, LLM_MODEL, LANGUAGE_INSTRUCTION
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

llm = ChatGroq(api_key=GROQ_API_KEY, model=LLM_MODEL, temperature=0.3)

def run_synthesis(results: dict, business_idea: str, location: str) -> str:
    
    language = results.get("language", "english")
    lang_instruction = LANGUAGE_INSTRUCTION.get(language, LANGUAGE_INSTRUCTION["english"])
    
    prompt = PromptTemplate(
        input_variables=["business_idea", "location", "market", "competitor",
                        "brand", "legal", "content", "financial", "lang_instruction"],
        template="""
You are an expert startup consultant.
IMPORTANT: {lang_instruction}

Business: {business_idea}
Location: {location}

Based on these reports:
MARKET RESEARCH: {market}
COMPETITOR ANALYSIS: {competitor}
BRAND IDENTITY: {brand}
LEGAL COMPLIANCE: {legal}
CONTENT STRATEGY: {content}
FINANCIAL MODEL: {financial}

Write a powerful Executive Summary covering:
1. Business Opportunity (2-3 sentences)
2. Target Market (2-3 sentences)
3. Competitive Advantage (2-3 sentences)
4. Recommended Brand Name & Tagline
5. Financial Highlights (investment needed, break-even, 6-month revenue)
6. Top 3 Action Items to start immediately

Be concise, powerful and specific. This is the first page of the startup package.
"""
    )
    
    chain = prompt | llm
    result = chain.invoke({
        "business_idea": business_idea,
        "location": location,
        "market": results["market_research"][:1000],
        "competitor": results["competitor_analysis"][:1000],
        "brand": results["brand_identity"][:1000],
        "legal": results["legal_compliance"][:1000],
        "content": results["content_generation"][:1000],
        "financial": results["financial_modeling"][:1000],
        "lang_instruction": lang_instruction
    })
    
    return result.content

def save_to_docx(results: dict, executive_summary: str,
                  business_idea: str, location: str) -> str:
    
    doc = Document()
    
    title = doc.add_heading('AutoFounder Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f'{business_idea} — {location}', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(executive_summary)
    doc.add_page_break()
    
    sections = {
        "Market Research": results["market_research"],
        "Competitor Analysis": results["competitor_analysis"],
        "Brand Identity": results["brand_identity"],
        "Legal Compliance": results["legal_compliance"],
        "Content Strategy": results["content_generation"],
        "Financial Model": results["financial_modeling"],
    }
    
    for section_title, content in sections.items():
        doc.add_heading(section_title, 1)
        doc.add_paragraph(content)
        doc.add_page_break()
    
    filename = f"output/AutoFounder_{business_idea.replace(' ', '_')}_{location}.docx"
    doc.save(filename)
    
    return filename